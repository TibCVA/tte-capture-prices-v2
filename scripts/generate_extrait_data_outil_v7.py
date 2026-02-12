from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
import sys
from typing import Any

import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.config_loader import load_countries


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate Extrait Data Outil v7 (docx + markdown) with audited outputs and test ledger."
    )
    parser.add_argument("--run-id", required=True, help="Combined run id under outputs/combined/<run_id>.")
    parser.add_argument("--countries", default="DE,ES,FR", help="Comma-separated country scope.")
    parser.add_argument("--output-docx", default="reports/Extrait Data Outil v7.docx", help="DOCX output path.")
    parser.add_argument(
        "--output-md",
        default="reports/Extrait Data Outil v7.md",
        help="Markdown output path (same content, audit-friendly).",
    )
    return parser.parse_args()


def _read_csv(path: Path, label: str) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"{label} missing: {path}")
    return pd.read_csv(path)


def _to_num(s: pd.Series) -> pd.Series:
    return pd.to_numeric(s, errors="coerce")


def _ratio(num: pd.Series, den: pd.Series) -> pd.Series:
    den_num = _to_num(den)
    num_num = _to_num(num)
    out = pd.Series(np.nan, index=num_num.index, dtype=float)
    mask = den_num > 0
    out.loc[mask] = num_num.loc[mask] / den_num.loc[mask]
    return out


def _country_timezones() -> dict[str, str]:
    cfg = load_countries()
    out: dict[str, str] = {}
    for country, params in cfg.get("countries", {}).items():
        out[str(country)] = str(params.get("timezone", "UTC"))
    return out


def _assumption_value(df: pd.DataFrame, param_name: str, default: float = np.nan) -> float:
    if df.empty:
        return default
    m = df["param_name"].astype(str) == str(param_name)
    if not m.any():
        return default
    try:
        return float(pd.to_numeric(df.loc[m, "param_value"], errors="coerce").iloc[0])
    except Exception:
        return default


def _pick(df: pd.DataFrame, *names: str, default: Any = np.nan) -> pd.Series:
    for name in names:
        if name in df.columns:
            return df[name]
    return pd.Series(default, index=df.index)


def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out = out.replace(r"^\s*$", np.nan, regex=True)
    return out


def _extract_annual_metrics_phase1(run_dir: Path, countries: list[str], tz_map: dict[str, str]) -> pd.DataFrame:
    df = _read_csv(run_dir / "Q1" / "hist" / "tables" / "Q1_year_panel.csv", "Q1_year_panel")
    df = df[df["country"].astype(str).isin(countries)].copy()
    if df.empty:
        return df
    if "scenario_id" not in df.columns:
        df["scenario_id"] = "HIST"
    df["scenario_id"] = df["scenario_id"].astype(str).replace({"": "HIST", "nan": "HIST", "None": "HIST"})

    out = pd.DataFrame(
        {
            "country": _pick(df, "country").astype(str),
            "year": _to_num(_pick(df, "year")).astype("Int64"),
            "scenario_id": _pick(df, "scenario_id").astype(str),
            "n_hours": _to_num(_pick(df, "n_hours_expected")).astype("Int64"),
            "timezone": _pick(df, "country").astype(str).map(tz_map).fillna("UTC"),
            "coverage_price": _to_num(_pick(df, "coverage_price")),
            "coverage_load_total": _to_num(_pick(df, "coverage_load_total")),
            "coverage_psh_pumping": _to_num(_pick(df, "coverage_psh_pumping", "psh_pumping_coverage_share")),
            "coverage_pv": _to_num(_pick(df, "coverage_pv")),
            "coverage_wind": _to_num(_pick(df, "coverage_wind")),
            "coverage_net_position": _to_num(_pick(df, "coverage_net_position")),
            "coverage_nuclear": _to_num(_pick(df, "coverage_nuclear")),
            "coverage_biomass": _to_num(_pick(df, "coverage_biomass")),
            "coverage_ror": _to_num(_pick(df, "coverage_ror")),
            "load_total_mw_avg": _to_num(_pick(df, "load_total_mw_avg")),
            "psh_pumping_mw_avg": _to_num(_pick(df, "psh_pumping_mw_avg")),
            "load_mw_avg": _to_num(_pick(df, "load_mw_avg")),
            "must_run_mw_avg": _to_num(_pick(df, "must_run_mw_avg")),
            "must_run_nuclear_mw_avg": _to_num(_pick(df, "must_run_nuclear_mw_avg")),
            "must_run_biomass_mw_avg": _to_num(_pick(df, "must_run_biomass_mw_avg")),
            "must_run_ror_mw_avg": _to_num(_pick(df, "must_run_ror_mw_avg")),
            "nrl_mw_avg": _to_num(_pick(df, "nrl_mw_avg")),
            "nrl_p10_mw": _to_num(_pick(df, "nrl_p10_mw")),
            "nrl_p50_mw": _to_num(_pick(df, "nrl_p50_mw")),
            "nrl_p90_mw": _to_num(_pick(df, "nrl_p90_mw")),
            "surplus_mwh_total": _to_num(_pick(df, "surplus_energy_twh", "surplus_twh")) * 1e6,
            "sr_energy": _to_num(_pick(df, "sr_energy")),
            "sr_hours": _to_num(_pick(df, "sr_hours", "sr_hours_share")),
            "far_energy": _to_num(_pick(df, "far_energy", "far_observed")),
            "far_hours": _to_num(_pick(df, "far_hours")),
            "sink_breakdown_json": _pick(df, "sink_breakdown_json"),
            "ir_p10": _to_num(_pick(df, "ir_p10")),
            "baseload_price_eur_mwh": _to_num(_pick(df, "baseload_price_eur_mwh")),
            "ttl_observed_eur_mwh": _to_num(_pick(df, "ttl_observed_eur_mwh", "ttl_eur_mwh")),
            "capture_price_pv_eur_mwh": _to_num(_pick(df, "capture_price_pv_eur_mwh", "capture_rate_pv_eur_mwh")),
            "capture_ratio_pv": _to_num(_pick(df, "capture_ratio_pv")),
            "capture_ratio_pv_vs_ttl_observed": _to_num(_pick(df, "capture_ratio_pv_vs_ttl_observed", "capture_ratio_pv_vs_ttl")),
            "capture_price_wind_eur_mwh": _to_num(_pick(df, "capture_price_wind_eur_mwh", "capture_rate_wind_eur_mwh")),
            "capture_ratio_wind": _to_num(_pick(df, "capture_ratio_wind")),
            "capture_ratio_wind_vs_ttl_observed": _to_num(_pick(df, "capture_ratio_wind_vs_ttl_observed", "capture_ratio_wind_vs_ttl")),
            "h_negative_obs": _to_num(_pick(df, "h_negative_obs")),
            "h_below_5_obs": _to_num(_pick(df, "h_below_5_obs")),
            "days_spread_gt50": _to_num(_pick(df, "days_spread_gt50", "days_spread_50_obs")),
            "regime_coherence": _to_num(_pick(df, "regime_coherence")),
            "nrl_price_corr": _to_num(_pick(df, "nrl_price_corr")),
            "load_identity_abs_max_mw": _to_num(_pick(df, "load_identity_abs_max_mw")),
            "load_identity_rel_err": _to_num(_pick(df, "load_identity_rel_err")),
            "load_identity_ok": _pick(df, "load_identity_ok"),
            "data_quality_flags": _pick(df, "data_quality_flags"),
            "quality_flag": _pick(df, "quality_flag"),
            "quality_notes": _pick(df, "warnings_quality", "core_sanity_issues"),
        }
    )
    out = _clean_df(out)
    out = out.drop_duplicates(subset=["country", "year", "scenario_id"], keep="first")
    out = out.sort_values(["country", "year", "scenario_id"]).reset_index(drop=True)
    return out


def _extract_q1_transition_summary(run_dir: Path, countries: list[str], phase1_assumptions: pd.DataFrame) -> pd.DataFrame:
    df = _read_csv(run_dir / "Q1" / "hist" / "tables" / "Q1_country_summary.csv", "Q1_country_summary")
    df = df[df["country"].astype(str).isin(countries)].copy()
    if df.empty:
        return df

    persistence = _assumption_value(phase1_assumptions, "q1_persistence_window_years", default=np.nan)

    def _stage(row: pd.Series) -> str:
        bm = pd.to_numeric(pd.Series([row.get("bascule_year_market")]), errors="coerce").iloc[0]
        if pd.notna(bm):
            return "2"
        s1 = str(row.get("stage1_detected", "")).strip().lower()
        if s1 in {"true", "1"}:
            return "1"
        return "NA"

    out = pd.DataFrame(
        {
            "country": df["country"].astype(str),
            "transition_year": pd.to_numeric(df.get("bascule_year_market"), errors="coerce").astype("Int64"),
            "stage": df.apply(_stage, axis=1),
            "stage2_score": pd.to_numeric(df.get("stage2_market_score_at_bascule"), errors="coerce"),
            "non_capture_flags_count": (
                pd.to_numeric(df.get("low_price_flags_count_at_bascule"), errors="coerce").fillna(0.0)
                + pd.to_numeric(df.get("physical_flags_count_at_bascule"), errors="coerce").fillna(0.0)
            ).astype("Int64"),
            "reason_codes": df.get("drivers_at_bascule", pd.Series("", index=df.index)).fillna("").astype(str),
            "persistence_window_years": persistence,
            "confidence_level": pd.to_numeric(df.get("bascule_confidence"), errors="coerce"),
        }
    )
    out = out.sort_values(["country"]).reset_index(drop=True)
    return out


def _extract_q2_slope_summary(run_dir: Path, countries: list[str]) -> pd.DataFrame:
    df = _load_table_across_modes(run_dir, "Q2", "Q2_country_slopes")
    if df.empty:
        df = _read_csv(run_dir / "Q2" / "hist" / "tables" / "Q2_country_slopes.csv", "Q2_country_slopes")
    df = df[df["country"].astype(str).isin(countries)].copy()
    if df.empty:
        return df
    if "scenario_id" not in df.columns:
        df["scenario_id"] = "HIST"
    df["scenario_id"] = _pick(df, "scenario_id", default="HIST").astype(str).replace({"": "HIST", "nan": "HIST", "None": "HIST"})

    def _driver_stats(row: pd.Series) -> str:
        payload = {
            "sr_energy_mean": pd.to_numeric(pd.Series([row.get("mean_sr_energy_phase2")]), errors="coerce").iloc[0],
            "far_energy_mean": pd.to_numeric(pd.Series([row.get("mean_far_energy_phase2")]), errors="coerce").iloc[0],
            "ir_p10_mean": pd.to_numeric(pd.Series([row.get("mean_ir_p10_phase2")]), errors="coerce").iloc[0],
            "ttl_mean": pd.to_numeric(pd.Series([row.get("mean_ttl_phase2")]), errors="coerce").iloc[0],
            "corr_vre_load_mean": pd.to_numeric(pd.Series([row.get("corr_vre_load_phase2")]), errors="coerce").iloc[0],
        }
        return json.dumps(payload, ensure_ascii=False)

    out = pd.DataFrame(
        {
            "scenario_id": df["scenario_id"].astype(str),
            "country": df["country"].astype(str),
            "tech": df["tech"].astype(str),
            "x_var_used": df.get("x_axis_used", pd.Series("", index=df.index)).astype(str),
            "y_var_used": "capture_ratio_" + df["tech"].astype(str),
            "n_points": pd.to_numeric(df.get("n_points", df.get("n", pd.Series(np.nan, index=df.index))), errors="coerce").astype("Int64"),
            "years_used": df.get("years_used", pd.Series("", index=df.index)).astype(str),
            "slope": pd.to_numeric(df.get("slope"), errors="coerce"),
            "intercept": pd.to_numeric(df.get("intercept"), errors="coerce"),
            "r2": pd.to_numeric(df.get("r2"), errors="coerce"),
            "p_value": pd.to_numeric(df.get("p_value"), errors="coerce"),
            "driver_stats_phase2": df.apply(_driver_stats, axis=1),
            "slope_quality_flag": df.get("slope_quality_flag", pd.Series("", index=df.index)).astype(str),
            "slope_quality_notes": df.get("slope_quality_notes", pd.Series("", index=df.index)).astype(str),
        }
    )
    out = out.sort_values(["country", "scenario_id", "tech"]).reset_index(drop=True)
    return out


def _extract_q3_inversion_requirements(run_dir: Path, countries: list[str], phase1_assumptions: pd.DataFrame) -> pd.DataFrame:
    preferred = run_dir / "Q3" / "hist" / "tables" / "q3_inversion_requirements.csv"
    if preferred.exists():
        df = _read_csv(preferred, "q3_inversion_requirements")
    else:
        df = _read_csv(run_dir / "Q3" / "hist" / "tables" / "Q3_status.csv", "Q3_status")
    df = df[df["country"].astype(str).isin(countries)].copy()
    if df.empty:
        return df

    if "lever" not in df.columns:
        target_sr = _assumption_value(phase1_assumptions, "stage1_sr_energy_max", default=np.nan)
        target_h_negative = _assumption_value(phase1_assumptions, "stage1_h_negative_max", default=np.nan)
        target_h_below_5 = _assumption_value(phase1_assumptions, "stage1_h_below_5_max", default=np.nan)
        rows: list[dict[str, Any]] = []
        for _, row in df.iterrows():
            country = str(row.get("country", ""))
            scenario_id = str(row.get("scenario_id", "HIST") or "HIST")
            year = _to_num(pd.Series([row.get("reference_year")])).iloc[0]
            rows.extend(
                [
                    {
                        "country": country,
                        "scenario_id": scenario_id,
                        "year": year,
                        "lever": "demand_uplift",
                        "required_uplift": _to_num(pd.Series([row.get("inversion_k_demand")])).iloc[0],
                        "required_uplift_mw": _to_num(pd.Series([row.get("required_uplift_mw", row.get("inversion_k_demand"))])).iloc[0],
                        "required_uplift_pct_avg_load": _to_num(pd.Series([row.get("required_uplift_pct_avg_load")])).iloc[0],
                        "required_uplift_twh_per_year": _to_num(pd.Series([row.get("required_uplift_twh_per_year")])).iloc[0],
                        "within_bounds": bool(row.get("within_bounds", False)),
                        "target_sr": target_sr,
                        "target_h_negative": target_h_negative,
                        "target_h_below_5": target_h_below_5,
                        "predicted_sr_after": _to_num(pd.Series([row.get("predicted_sr_after")])).iloc[0],
                        "predicted_far_after": _to_num(pd.Series([row.get("predicted_far_after")])).iloc[0],
                        "predicted_h_negative_after": _to_num(pd.Series([row.get("predicted_h_negative_after")])).iloc[0],
                        "predicted_h_below_5_after": _to_num(pd.Series([row.get("predicted_h_below_5_after")])).iloc[0],
                        "predicted_h_negative_metric": str(row.get("predicted_h_negative_metric", "MARKET_PROXY_NRL_LOOKUP")),
                        "applicability_flag": "APPLICABLE" if bool(row.get("in_phase2", False)) else "HORS_SCOPE_PHASE2",
                        "status": str(row.get("status", "missing_data")),
                        "reason": str(row.get("reason_code", "")),
                    },
                    {
                        "country": country,
                        "scenario_id": scenario_id,
                        "year": year,
                        "lever": "export_uplift",
                        "required_uplift": np.nan,
                        "required_uplift_mw": np.nan,
                        "required_uplift_pct_avg_load": np.nan,
                        "required_uplift_twh_per_year": np.nan,
                        "within_bounds": False,
                        "target_sr": target_sr,
                        "target_h_negative": target_h_negative,
                        "target_h_below_5": target_h_below_5,
                        "predicted_sr_after": np.nan,
                        "predicted_far_after": np.nan,
                        "predicted_h_negative_after": np.nan,
                        "predicted_h_below_5_after": np.nan,
                        "predicted_h_negative_metric": "MARKET_PROXY_NRL_LOOKUP",
                        "applicability_flag": "APPLICABLE" if bool(row.get("in_phase2", False)) else "HORS_SCOPE_PHASE2",
                        "status": "missing_data",
                        "reason": "legacy_q3_status_no_export_solver",
                    },
                    {
                        "country": country,
                        "scenario_id": scenario_id,
                        "year": year,
                        "lever": "flex_uplift",
                        "required_uplift": _to_num(pd.Series([row.get("inversion_f_flex")])).iloc[0],
                        "required_uplift_mw": _to_num(pd.Series([row.get("required_uplift_mw", row.get("inversion_f_flex"))])).iloc[0],
                        "required_uplift_pct_avg_load": _to_num(pd.Series([row.get("required_uplift_pct_avg_load")])).iloc[0],
                        "required_uplift_twh_per_year": _to_num(pd.Series([row.get("required_uplift_twh_per_year")])).iloc[0],
                        "within_bounds": False,
                        "target_sr": target_sr,
                        "target_h_negative": target_h_negative,
                        "target_h_below_5": target_h_below_5,
                        "predicted_sr_after": np.nan,
                        "predicted_far_after": np.nan,
                        "predicted_h_negative_after": np.nan,
                        "predicted_h_below_5_after": np.nan,
                        "predicted_h_negative_metric": "MARKET_PROXY_NRL_LOOKUP",
                        "applicability_flag": "APPLICABLE" if bool(row.get("in_phase2", False)) else "HORS_SCOPE_PHASE2",
                        "status": str(row.get("inversion_f_flex_status", "missing_data")),
                        "reason": str(row.get("reason_code", "")),
                    },
                ]
            )
        df = pd.DataFrame(rows)

    for col, default in {
        "scenario_id": "HIST",
        "year": np.nan,
        "status": "missing_data",
        "reason": "",
        "predicted_h_negative_metric": "MARKET_PROXY_NRL_LOOKUP",
    }.items():
        if col not in df.columns:
            df[col] = default

    keep = [
        "country",
        "scenario_id",
        "year",
        "lever",
        "required_uplift",
        "required_uplift_mw",
        "required_uplift_pct_avg_load",
        "required_uplift_twh_per_year",
        "within_bounds",
        "target_sr",
        "target_h_negative",
        "target_h_below_5",
        "predicted_sr_after",
        "predicted_far_after",
        "predicted_h_negative_after",
        "predicted_h_below_5_after",
        "predicted_h_negative_metric",
        "applicability_flag",
        "status",
        "reason",
    ]
    out = _clean_df(df[[c for c in keep if c in df.columns]].copy())
    out = out.sort_values(["country", "scenario_id", "year", "lever"]).reset_index(drop=True)
    return out


def _load_table_across_modes(run_dir: Path, qid: str, table_name: str) -> pd.DataFrame:
    paths: list[Path] = [run_dir / qid / "hist" / "tables" / f"{table_name}.csv"]
    scen_root = run_dir / qid / "scen"
    if scen_root.exists():
        for scen_dir in sorted([p for p in scen_root.iterdir() if p.is_dir()]):
            paths.append(scen_dir / "tables" / f"{table_name}.csv")
    frames: list[pd.DataFrame] = []
    for p in paths:
        if p.exists():
            frames.append(pd.read_csv(p))
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def _extract_q4_bess_sizing_curve(run_dir: Path, countries: list[str]) -> pd.DataFrame:
    df = _load_table_across_modes(run_dir, "Q4", "q4_bess_sizing_curve")
    if df.empty:
        df = _load_table_across_modes(run_dir, "Q4", "Q4_bess_frontier")
    if df.empty:
        return df
    df = df[df["country"].astype(str).isin(countries)].copy()
    if df.empty:
        return df

    power_mw = pd.to_numeric(_pick(df, "bess_power_mw_test", "bess_power_mw"), errors="coerce")
    energy_mwh = pd.to_numeric(_pick(df, "bess_energy_mwh_test", "bess_energy_mwh"), errors="coerce")

    if "cycles_realized_per_day" in df.columns:
        cycles_realized = pd.to_numeric(df["cycles_realized_per_day"], errors="coerce")
    else:
        discharge_sum = pd.to_numeric(_pick(df, "discharge_sum_mwh"), errors="coerce")
        year_vals = pd.to_numeric(_pick(df, "year"), errors="coerce")
        n_days = pd.Series(
            [366.0 if pd.notna(y) and int(y) % 4 == 0 else 365.0 for y in year_vals],
            index=df.index,
            dtype=float,
        )
        cycles_realized = pd.Series(0.0, index=df.index, dtype=float)
        m = energy_mwh > 0
        cycles_realized.loc[m] = discharge_sum.loc[m] / (energy_mwh.loc[m] * n_days.loc[m])

    out = pd.DataFrame(
        {
            "country": _pick(df, "country").astype(str),
            "scenario_id": _pick(df, "scenario_id", default="HIST").astype(str).replace({"": "HIST", "nan": "HIST", "None": "HIST"}),
            "year": pd.to_numeric(_pick(df, "year"), errors="coerce").astype("Int64"),
            "bess_power_gw": power_mw / 1000.0,
            "bess_energy_gwh": energy_mwh / 1000.0,
            "cycles_realized_per_day": pd.to_numeric(cycles_realized, errors="coerce"),
            "eta_roundtrip": pd.to_numeric(_pick(df, "eta_roundtrip"), errors="coerce"),
            "far_energy_after": pd.to_numeric(_pick(df, "far_energy_after", "far_after"), errors="coerce"),
            "surplus_unabsorbed_twh_after": pd.to_numeric(_pick(df, "surplus_unabsorbed_twh_after", "surplus_unabs_energy_after"), errors="coerce"),
            "h_negative_proxy_after": pd.to_numeric(_pick(df, "h_negative_proxy_after", "h_negative_after"), errors="coerce"),
            "h_negative_reducible_upper_bound": pd.to_numeric(_pick(df, "h_negative_reducible_upper_bound"), errors="coerce"),
            "monotonicity_check_flag": _pick(df, "monotonicity_check_flag", default="PASS").astype(str),
            "physics_check_flag": _pick(df, "physics_check_flag", default="PASS").astype(str),
            "on_efficient_frontier": _pick(df, "on_efficient_frontier", default=False),
            "notes": (
                "dispatch_mode="
                + _pick(df, "dispatch_mode", default="").astype(str)
                + "; objective="
                + _pick(df, "objective", default="").astype(str)
            ),
        }
    )
    out = _clean_df(out)
    out = out.drop_duplicates(
        subset=["country", "scenario_id", "year", "bess_power_gw", "bess_energy_gwh", "eta_roundtrip"],
        keep="first",
    )
    out = out.sort_values(["country", "scenario_id", "year", "bess_power_gw", "bess_energy_gwh"]).reset_index(drop=True)
    return out


def _extract_q5_anchor_sensitivity(run_dir: Path, countries: list[str]) -> pd.DataFrame:
    df = _load_table_across_modes(run_dir, "Q5", "q5_anchor_sensitivity")
    if df.empty:
        df = _load_table_across_modes(run_dir, "Q5", "Q5_summary")
    if df.empty:
        return df
    df = df[df["country"].astype(str).isin(countries)].copy()
    if df.empty:
        return df

    out = pd.DataFrame(
        {
            "country": _pick(df, "country").astype(str),
            "scenario_id": _pick(df, "scenario_id", default="HIST").astype(str).replace({"": "HIST", "nan": "HIST", "None": "HIST"}),
            "year": pd.to_numeric(_pick(df, "year", "ttl_reference_year"), errors="coerce").astype("Int64"),
            "gas_eur_per_mwh_th": pd.to_numeric(_pick(df, "gas_eur_per_mwh_th", "assumed_gas_price_eur_mwh_th"), errors="coerce"),
            "co2_eur_per_t": pd.to_numeric(_pick(df, "co2_eur_per_t", "assumed_co2_price_eur_t"), errors="coerce"),
            "tca_ccgt_eur_mwh": pd.to_numeric(_pick(df, "tca_ccgt_eur_mwh", "tca_q95"), errors="coerce"),
            "tca_coal_eur_mwh": pd.to_numeric(_pick(df, "tca_coal_eur_mwh"), errors="coerce"),
            "ttl_observed_eur_mwh": pd.to_numeric(_pick(df, "ttl_observed_eur_mwh", "ttl_obs"), errors="coerce"),
            "ttl_model_eur_mwh": pd.to_numeric(_pick(df, "ttl_model_eur_mwh", "ttl_anchor_formula", "ttl_anchor"), errors="coerce"),
            "delta_tca_vs_base": pd.to_numeric(_pick(df, "delta_tca_vs_base"), errors="coerce"),
            "delta_ttl_model_vs_base": pd.to_numeric(_pick(df, "delta_ttl_model_vs_base"), errors="coerce"),
            "coherence_flag": _pick(df, "coherence_flag", default="").astype(str),
            "ttl_proxy_method": _pick(df, "ttl_proxy_method", default="").astype(str),
            "status": _pick(df, "status", default="").astype(str),
            "reason": _pick(df, "reason", default="").astype(str),
            "base_ref_status": _pick(df, "base_ref_status", default="").astype(str),
            "base_ref_reason": _pick(df, "base_ref_reason", default="").astype(str),
        }
    )

    base = out[out["scenario_id"].astype(str).str.upper() == "BASE"][
        ["country", "year", "tca_ccgt_eur_mwh", "ttl_model_eur_mwh"]
    ].rename(
        columns={
            "tca_ccgt_eur_mwh": "base_tca",
            "ttl_model_eur_mwh": "base_ttl_model",
        }
    )
    out = out.merge(base, on=["country", "year"], how="left")
    missing_base_mask = out["status"].astype(str).str.lower().eq("missing_base") | out["base_ref_status"].astype(str).str.lower().eq("missing_base")
    out["delta_tca_vs_base"] = out["delta_tca_vs_base"].where(
        out["delta_tca_vs_base"].notna() | missing_base_mask,
        out["tca_ccgt_eur_mwh"] - out["base_tca"],
    )
    out["delta_ttl_model_vs_base"] = out["delta_ttl_model_vs_base"].where(
        out["delta_ttl_model_vs_base"].notna() | missing_base_mask,
        out["ttl_model_eur_mwh"] - out["base_ttl_model"],
    )
    hist_mask = out["scenario_id"].astype(str).str.upper() == "HIST"
    out.loc[hist_mask, "delta_tca_vs_base"] = out.loc[hist_mask, "delta_tca_vs_base"].fillna(0.0)
    out.loc[hist_mask, "delta_ttl_model_vs_base"] = out.loc[hist_mask, "delta_ttl_model_vs_base"].fillna(0.0)

    if (out["coherence_flag"].astype(str).str.strip() == "").any():
        sign_ok = (
            np.sign(out["delta_tca_vs_base"].fillna(0.0))
            == np.sign(out["delta_ttl_model_vs_base"].fillna(0.0))
        )
        out["coherence_flag"] = np.where(sign_ok, "PASS", "WARN")

    out = out.drop(columns=["base_tca", "base_ttl_model"], errors="ignore")
    out = _clean_df(out)
    out = out.sort_values(["country", "scenario_id", "year"]).reset_index(drop=True)
    return out


def _extract_country_year(msg: str) -> tuple[str, str]:
    text = str(msg or "")
    m = re.search(r"\b([A-Z_]{2,})[- ](\d{4})\b", text)
    if m:
        return m.group(1), m.group(2)
    return "", ""


def _build_test_ledger(run_dir: Path, countries: list[str]) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for qid in ["Q1", "Q2", "Q3", "Q4", "Q5"]:
        tl_path = run_dir / qid / "test_ledger.csv"
        if tl_path.exists():
            tl = pd.read_csv(tl_path)
            for _, r in tl.iterrows():
                rows.append(
                    {
                        "test_id": str(r.get("test_id", "")),
                        "scope": str(r.get("question_id", qid)),
                        "country": np.nan,
                        "year": np.nan,
                        "scenario_id": str(r.get("scenario_id", "")),
                        "status": str(r.get("status", "")),
                        "metric_name": str(r.get("title", "")),
                        "observed_value": str(r.get("value", np.nan)),
                        "expected_rule": str(r.get("metric_rule", np.nan)),
                        "message": str(r.get("interpretation", "")),
                    }
                )

        summary_path = run_dir / qid / "summary.json"
        if summary_path.exists():
            summary = json.loads(summary_path.read_text(encoding="utf-8"))
            for c in summary.get("checks", []):
                msg = str(c.get("message", ""))
                country, year = _extract_country_year(msg)
                if country and country not in countries:
                    continue
                rows.append(
                    {
                        "test_id": str(c.get("code", "")),
                        "scope": qid,
                        "country": country,
                        "year": year,
                        "scenario_id": str(c.get("scenario_id", "")),
                        "status": str(c.get("status", "")),
                        "metric_name": str(c.get("code", "")),
                        "observed_value": np.nan,
                        "expected_rule": np.nan,
                        "message": msg,
                    }
                )

    ledger = _clean_df(pd.DataFrame(rows))
    if not ledger.empty:
        ledger = ledger.sort_values(["scope", "test_id", "country", "year", "scenario_id"]).reset_index(drop=True)
    return ledger


def _df_csv_block(df: pd.DataFrame) -> str:
    clean = _clean_df(df)
    return f"```csv\n{clean.to_csv(index=False, na_rep='NaN').rstrip()}\n```"


def _assert_non_empty(outputs: dict[str, pd.DataFrame]) -> None:
    for name, df in outputs.items():
        if df.empty:
            raise ValueError(f"{name} is empty or not computed; export aborted.")


def _add_csv_block_to_doc(doc: Any, title: str, df: pd.DataFrame) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph("```csv")
    csv_text = _clean_df(df).to_csv(index=False, na_rep="NaN").rstrip().splitlines()
    for line in csv_text:
        doc.add_paragraph(line)
    doc.add_paragraph("```")


def generate_extrait(run_id: str, countries: list[str], output_docx: Path, output_md: Path) -> dict[str, Any]:
    run_dir = ROOT / "outputs" / "combined" / run_id
    if not run_dir.exists():
        raise FileNotFoundError(f"Run directory not found: {run_dir}")

    phase1_assumptions = _read_csv(ROOT / "data" / "assumptions" / "phase1_assumptions.csv", "phase1_assumptions")
    phase2_assumptions = _read_csv(
        ROOT / "data" / "assumptions" / "phase2" / "phase2_scenario_country_year.csv",
        "phase2_scenario_country_year",
    )
    phase2_scope = phase2_assumptions[phase2_assumptions["country"].astype(str).isin(countries)].copy()
    tz_map = _country_timezones()

    annual_metrics_phase1 = _extract_annual_metrics_phase1(run_dir, countries, tz_map)
    q1_transition_summary = _extract_q1_transition_summary(run_dir, countries, phase1_assumptions)
    q2_slope_summary = _extract_q2_slope_summary(run_dir, countries)
    q3_inversion_requirements = _extract_q3_inversion_requirements(run_dir, countries, phase1_assumptions)
    q4_bess_sizing_curve = _extract_q4_bess_sizing_curve(run_dir, countries)
    q5_anchor_sensitivity = _extract_q5_anchor_sensitivity(run_dir, countries)
    test_ledger = _build_test_ledger(run_dir, countries)

    outputs = {
        "annual_metrics_phase1": annual_metrics_phase1,
        "q1_transition_summary": q1_transition_summary,
        "q2_slope_summary": q2_slope_summary,
        "q3_inversion_requirements": q3_inversion_requirements,
        "q4_bess_sizing_curve": q4_bess_sizing_curve,
        "q5_anchor_sensitivity": q5_anchor_sensitivity,
        "test_ledger": test_ledger,
    }
    _assert_non_empty(outputs)

    q1_summary = json.loads((run_dir / "Q1" / "summary.json").read_text(encoding="utf-8"))
    meta = {
        "run_id": run_id,
        "run_dir": str(run_dir),
        "countries": countries,
        "hist_years": q1_summary.get("selection", {}).get("years", []),
        "scenario_years": q1_summary.get("selection", {}).get("scenario_years", []),
        "scenarios": q1_summary.get("selection", {}).get("scenario_ids", []),
        "method_reference": "AUDIT_METHODS_Q1_Q5.md",
    }

    lines: list[str] = []
    lines.append("# Extrait Data Outil v7")
    lines.append("")
    lines.append("## AUDIT PAYLOAD")
    lines.append("")
    lines.append("### Inputs")
    lines.append("")
    lines.append("#### run_metadata")
    lines.append("```json")
    lines.append(json.dumps(meta, ensure_ascii=False, indent=2))
    lines.append("```")
    lines.append("")
    lines.append("#### Method Reference")
    lines.append("```text")
    lines.append(str(ROOT / "AUDIT_METHODS_Q1_Q5.md"))
    lines.append("```")
    lines.append("")
    lines.append("#### parameter_catalog")
    lines.append(_df_csv_block(phase1_assumptions))
    lines.append("")
    lines.append("#### Phase 2 assumptions")
    lines.append(_df_csv_block(phase2_scope))
    lines.append("")
    lines.append("## RUN OUTPUTS")
    lines.append("")
    lines.append("### A.1 annual_metrics_phase1")
    lines.append(_df_csv_block(annual_metrics_phase1))
    lines.append("")
    lines.append("### A.2 q1_transition_summary")
    lines.append(_df_csv_block(q1_transition_summary))
    lines.append("")
    lines.append("### A.3 q2_slope_summary")
    lines.append(_df_csv_block(q2_slope_summary))
    lines.append("")
    lines.append("### A.4 q3_inversion_requirements")
    lines.append(_df_csv_block(q3_inversion_requirements))
    lines.append("")
    lines.append("### A.5 q4_bess_sizing_curve")
    lines.append(_df_csv_block(q4_bess_sizing_curve))
    lines.append("")
    lines.append("### A.6 q5_anchor_sensitivity")
    lines.append(_df_csv_block(q5_anchor_sensitivity))
    lines.append("")
    lines.append("## TEST LEDGER")
    lines.append("")
    lines.append(_df_csv_block(test_ledger))
    lines.append("")

    output_md.parent.mkdir(parents=True, exist_ok=True)
    output_md.write_text("\n".join(lines), encoding="utf-8")

    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"python-docx unavailable: {exc}") from exc

    doc = Document()
    doc.add_heading("Extrait Data Outil v7", level=0)
    doc.add_heading("AUDIT PAYLOAD", level=1)
    doc.add_heading("Inputs", level=2)
    doc.add_heading("run_metadata", level=3)
    doc.add_paragraph("```json")
    for line in json.dumps(meta, ensure_ascii=False, indent=2).splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph("```")

    doc.add_heading("Method Reference", level=3)
    doc.add_paragraph("```text")
    doc.add_paragraph(str(ROOT / "AUDIT_METHODS_Q1_Q5.md"))
    doc.add_paragraph("```")

    _add_csv_block_to_doc(doc, "parameter_catalog", phase1_assumptions)
    _add_csv_block_to_doc(doc, "Phase 2 assumptions", phase2_scope)

    doc.add_heading("RUN OUTPUTS", level=1)
    _add_csv_block_to_doc(doc, "A.1 annual_metrics_phase1", annual_metrics_phase1)
    _add_csv_block_to_doc(doc, "A.2 q1_transition_summary", q1_transition_summary)
    _add_csv_block_to_doc(doc, "A.3 q2_slope_summary", q2_slope_summary)
    _add_csv_block_to_doc(doc, "A.4 q3_inversion_requirements", q3_inversion_requirements)
    _add_csv_block_to_doc(doc, "A.5 q4_bess_sizing_curve", q4_bess_sizing_curve)
    _add_csv_block_to_doc(doc, "A.6 q5_anchor_sensitivity", q5_anchor_sensitivity)

    doc.add_heading("TEST LEDGER", level=1)
    doc.add_paragraph("```csv")
    for line in _clean_df(test_ledger).to_csv(index=False, na_rep="NaN").rstrip().splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph("```")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    return {
        "run_id": run_id,
        "countries": countries,
        "output_docx": str(output_docx),
        "output_md": str(output_md),
        "rows": {k: int(len(v)) for k, v in outputs.items()},
    }


def main() -> int:
    args = _parse_args()
    countries = [c.strip() for c in str(args.countries).split(",") if c.strip()]
    result = generate_extrait(
        run_id=str(args.run_id),
        countries=countries,
        output_docx=Path(args.output_docx),
        output_md=Path(args.output_md),
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
